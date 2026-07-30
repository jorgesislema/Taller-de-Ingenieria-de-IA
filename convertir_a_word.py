#!/usr/bin/env python3
"""
Script para convertir apuntes.md a formato Word profesional.
Corrige ortografía y aplica formato profesional.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document():
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Título principal
    title = doc.add_heading('Apuntes de la Clase 5', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Arquitectura, Estructura y Mentalidad de Ingeniería para IA', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    # Introducción
    doc.add_heading('Introducción', level=2)
    
    intro_quote = doc.add_paragraph()
    intro_quote.style = doc.styles['Intense Quote'] if 'Intense Quote' in [s.name for s in doc.styles] else doc.styles['Normal']
    intro_quote.add_run('"Hasta ahora hemos aprendido a usar el martillo (VS Code) y el almacén (GitHub). Pero construir una casa no es solo golpear clavos. Hoy vamos a hablar de los planos. Hoy vamos a pensar como Ingenieros. Y un ingeniero no es el que sabe más código, es el que toma mejores decisiones basándose en el contexto."')
    
    doc.add_page_break()
    
    # Fase 1: El Cambio de Paradigma
    doc.add_heading('Fase 1: El Cambio de Paradigma — De "Enciclopedias" a "Directores de Orquesta"', level=2)
    
    doc.add_heading('La muerte del trabajador-enciclopedia', level=3)
    
    doc.add_paragraph('Durante los últimos 30 años, el trabajador del conocimiento fue pagado por su capacidad de almacenar y recordar. Si sabías la sintaxis exacta de 15 librerías de Python, o te sabías de memoria los protocolos de red, eras valioso. Eras una enciclopedia con piernas.')
    
    p = doc.add_paragraph()
    run = p.add_run('Eso murió.')
    run.bold = True
    
    doc.add_paragraph('ChatGPT, Claude, DeepSeek y los modelos que vengan tienen absorbida toda la enciclopedia de la humanidad. En un concurso de "quién sabe más datos", un humano pierde 100 a 0 en un segundo.')
    
    doc.add_paragraph('Entonces, ¿por qué una empresa te va a pagar un buen salario hoy? Te pagará por tres cosas que la IA no tiene:')
    
    # Las 3 cosas que la IA no puede hacer
    doc.add_heading('Las 3 cosas que la IA no puede hacer', level=3)
    
    doc.add_heading('1. Contexto del Negocio Real', level=4)
    doc.add_paragraph('La IA no sabe que tu jefe es terco, que el presupuesto es de 50 dólares, o que el cliente final odia los colores brillantes. Tú sí.')
    doc.add_paragraph('La IA puede generar 10 soluciones técnicamente perfectas, pero no sabe cuál se ajusta a la política interna de tu empresa, a las limitaciones presupuestarias del cliente, o al hecho de que el equipo de soporte solo sabe usar Excel.')
    
    doc.add_heading('2. Criterio bajo Incertidumbre', level=4)
    doc.add_paragraph('La IA te da 5 opciones técnicas válidas. ¿Cuál eliges? Depende de factores que no están en un manual. Eso es criterio.')
    doc.add_paragraph('El criterio es lo que te permite decir: "Sé que los microservicios son la moda, pero para nuestro equipo de 3 personas, un Monolito es la decisión correcta". La IA nunca te dirá eso porque no tiene el contexto de tu realidad.')
    
    doc.add_heading('3. Asumir las Consecuencias', level=4)
    doc.add_paragraph('Si la IA sugiere borrar una base de datos y lo hace, la IA no va a la cárcel ni es despedida. Tú sí. Alguien humano tiene que firmar y ser responsable.')
    doc.add_paragraph('La IA puede sugerir, pero el humano firma con su nombre. Por eso te pagan: no por saber, sino por responsabilidad.')
    
    # Ejemplos de Criterio vs. Conocimiento
    doc.add_heading('Ejemplos de Criterio vs. Conocimiento (Aplicados a esta clase)', level=3)
    
    doc.add_paragraph('Si leemos la clase que armamos a través de este lente, el "conocimiento" es solo el 10%. El 90% restante es puro criterio arquitectónico:')
    
    # Ejemplo 1
    doc.add_heading('Ejemplo 1: Las Arquitecturas (Monolito vs. Microservicios)', level=4)
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = 'El Conocimiento (Lo que sabe la IA)'
    table.cell(0, 1).text = 'El Criterio (Lo que pagamos a ti)'
    table.cell(1, 0).text = 'La IA sabe perfectamente definir qué es un microservicio, cómo se configura Kubernetes y cómo hacer un API Gateway.'
    table.cell(1, 1).text = 'Darse cuenta de que, para un equipo de 3 personas haciendo un MVP, usar microservicios es un suicidio financiero y operativo, y decidir valientemente usar un Monolito (Food Truck) a pesar de que la "moda" diga lo contrario. Eso es pensamiento crítico.'
    
    # Ejemplo 2
    doc.add_heading('Ejemplo 2: La regla del snake_case', level=4)
    
    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'
    
    table2.cell(0, 0).text = 'El Conocimiento'
    table2.cell(0, 1).text = 'El Criterio'
    table2.cell(1, 0).text = 'Saber que en Linux los espacios rompen las rutas.'
    table2.cell(1, 1).text = 'Exigirle a la IA a través de un RULES.md que use snake_case antes de que empiece a programar, para no tener que perder 2 horas buscando por qué el servidor se cayó por un error tonto. Eso es previsión.'
    
    # Ejemplo 3
    doc.add_heading('Ejemplo 3: El equipo Híbrido (DeepSeek para crear, Claude para auditar)', level=4)
    
    table3 = doc.add_table(rows=2, cols=2)
    table3.style = 'Table Grid'
    
    table3.cell(0, 0).text = 'El Conocimiento'
    table3.cell(0, 1).text = 'El Criterio'
    table3.cell(1, 0).text = 'Saber que existen dos modelos de IA distintos.'
    table3.cell(1, 1).text = 'Diseñar el flujo de trabajo de "Red Team". Entender que los modelos tienen "puntos ciegos" y que usar dos cerebros artificiales distintos (uno para atacar, otro para defender) genera un resultado superior al de un solo modelo. Eso es diseño de sistemas.'
    
    # La Analogía del Capitán del Barco
    doc.add_heading('La Analogía del Capitán del Barco', level=3)
    
    quote = doc.add_paragraph()
    quote.style = doc.styles['Intense Quote'] if 'Intense Quote' in [s.name for s in doc.styles] else doc.styles['Normal']
    quote.add_run('"Imaginen un barco carguero gigante. La IA es la sala de máquinas. Tiene una fuerza bruta increíble, sabe cómo mover cada válvula, conoce la termodinámica perfecta y no duerme. Tiene todo el conocimiento físico del barco.\n\nPero si la sala de máquinas decide hacia dónde ir el barco, va a chocar contra la primera isla que encuentre.\n\nEl trabajador del futuro no es el que está abajo shoveling carbón (escribiendo código repetitivo). El trabajador del futuro es el Capitán en el puente de mando. El Capitán no sabe cómo funcionan los tornillos del motor, pero sabe leer el clima, sabe a qué puerto quiere llegar, sabe negociar con otros barcos y, sobre todo, tiene el criterio para decidir si debe ir a toda velocidad o detenerse por la tormenta."')
    
    # Conclusión de la Analogía
    doc.add_heading('Conclusión de la Analogía', level=3)
    
    doc.add_paragraph('Lo que les estás enseñando en esta clase no es a "usar herramientas de IA". Les estás enseñando a subir al puente de mando.')
    doc.add_paragraph('El .md, las carpetas, las decisiones de VPS vs Serverless... no son archivos técnicos. Son los instrumentos de navegación del capitán. Y por eso, el que sabe usarlos, será imparable y altamente cotizado, sin importar si sabe escribir una sola línea de código.')
    
    # El Criterio como Multiplicador
    doc.add_heading('El Criterio como Multiplicador', level=3)
    
    doc.add_paragraph('El criterio no es una línea recta que sube con el tiempo. Es un multiplicador. Y la fórmula es esta:')
    
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run('Criterio = (Cantidad de Errores Comprendidos) × (Densidad de Reflexión)')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph('Fíjate que en la fórmula no aparece el tiempo. Puedes tener 20 años de experiencia y no haber cometido un error porque nunca te arriesgaste, o porque siempre hiciste lo que te dijeron. Tu criterio será de cero.')
    
    # Cómo Aumentar tu Criterio
    doc.add_heading('Cómo Aumentar tu Criterio (Sin Importar si Eres Principiante o Veterano)', level=3)
    
    doc.add_heading('1. El Secreto: El "Fracaso Digerido" (No basta con fallar)', level=4)
    doc.add_paragraph('El éxito no te da criterio. El éxito te da confianza (que a veces es peligrosa). El criterio nace exclusivamente del dolor de haberse equivocado y haber entendido por qué.')
    
    table4 = doc.add_table(rows=2, cols=2)
    table4.style = 'Table Grid'
    
    table4.cell(0, 0).text = 'Sin experiencia'
    table4.cell(0, 1).text = 'Construyendo criterio'
    table4.cell(1, 0).text = 'Si hoy creas tu archivo .md, se lo das a la IA, el código falla, te frustras y lo borras... no ganaste criterio.'
    table4.cell(1, 1).text = 'Si el código falla, te detienes, analizas y dices: "Ah, falló porque le di a la IA un CONTEXT.md muy vago y se inventó una librería que no existe". Acabas de ganar +10 puntos de criterio.'
    
    doc.add_paragraph('Cómo aplicarlo sin experiencia: Fuerza errores pequeños y baratos. En un entorno de pruebas, dile a la IA que haga algo mal a propósito solo para ver cómo se rompe el sistema. Estudiar el rompimiento construye criterio más rápido que estudiar el éxito.')
    
    doc.add_heading('2. La Densidad vs. La Longitud (Calidad de las decisiones)', level=4)
    doc.add_paragraph('Una persona que en un mes tiene que tomar 50 decisiones difíciles sobre arquitectura de IA (aunque se equivoque en 30), desarrollará mucho más criterio que alguien que en 5 años solo toma 2 decisiones al año porque trabaja en una empresa burocrática donde todo ya está decidido.')
    doc.add_paragraph('Cómo aplicarlo sin experiencia: Aumenta tu densidad. En lugar de pedirle a la IA "Hazme el proyecto entero", pártelo en 10 pedacitos. En cada pedacito, tú toma una decisión de diseño (¿Cómo llamo esta carpeta? ¿Pongo esto en un .md o en otro?). Toma 50 micro-decisiones al día, aunque al principio no sepas cuál es la correcta.')
    
    doc.add_heading('3. El Método del "Por Qué" en Cadena (El arte de incomodar)', level=4)
    doc.add_paragraph('El criterio se destruye cuando aceptamos las cosas porque "así se hace siempre".')
    doc.add_paragraph('Cómo aplicarlo sin experiencia: Conviértete en un niño de 3 años. Cuando veas un tutorial, un libro, o cuando la IA te dé una solución, aplica la regla de los 3 Porqués:')
    
    # Ejemplo de diálogo
    dialogo = doc.add_paragraph()
    dialogo.style = doc.styles['Intense Quote'] if 'Intense Quote' in [s.name for s in doc.styles] else doc.styles['Normal']
    dialogo.add_run('La IA dice: "Usa un archivo .env para la contraseña."\nTú: ¿Por qué?\nLa IA: "Para no subirla a GitHub."\nTú: ¿Y por qué es malo subirla a GitHub?\nLa IA: "Porque los bots escanean los repositorios públicos para robar claves y cobrar en AWS."\nTú: ¡Aha! Ya no es una regla memorizada, es un criterio de seguridad interno.')
    
    doc.add_heading('4. Desarrollar la "Simulación Mental" (El Gym del Cerebro)', level=4)
    doc.add_paragraph('Los grandes arquitectos de software y los grandes ajedrecistas no son más rápidos calculando; son mejores simulando el futuro en su cabeza antes de mover la pieza.')
    doc.add_paragraph('Cómo aplicarlo sin experiencia: Antes de escribirle a la IA o de crear una carpeta, siéntate 60 segundos. Cierra los ojos y visualiza:')
    
    quote2 = doc.add_paragraph()
    quote2.style = doc.styles['Intense Quote'] if 'Intense Quote' in [s.name for s in doc.styles] else doc.styles['Normal']
    quote2.add_run('"Si yo creo esta carpeta llamada archivos y mañana vienen 5 personas más al equipo... ¿van a saber qué va ahí? Si la IA genera 100 archivos ahí, ¿será fácil encontrar algo?"')
    
    doc.add_paragraph('Ese segundo de pausa, esa proyección al futuro, es criterio puro. Y no requiere saber programar, requiere imaginación estructural.')
    
    # La Ventaja del que "No Tiene Experiencia"
    doc.add_heading('La Ventaja del que "No Tiene Experiencia"', level=3)
    
    doc.add_paragraph('Tengo una buena noticia como docente: El que no tiene experiencia tiene una ventaja gigantesca sobre el veterano: NO tiene prejuicios.')
    doc.add_paragraph('El programador con 15 años de experiencia va a querer hacer las cosas como se hacían en 2015. Va a resistirse a usar archivos .md para controlar la IA porque "eso no es programar de verdad". Su conocimiento viejo bloquea su criterio nuevo.')
    doc.add_paragraph('Tú, al empezar desde cero con la IA, tienes la mente en blanco. Tu cerebro es plastilina fresca. Si te concentras en tomar decisiones conscientes (aunque duelen), analizar tus errores y no memorizar recetas, puedes alcanzar en 6 meses un nivel de criterio arquitectónico que a un tradicional le tomó 4 años.')
    
    quote3 = doc.add_paragraph()
    quote3.style = doc.styles['Intense Quote'] if 'Intense Quote' in [s.name for s in doc.styles] else doc.styles['Normal']
    quote3.add_run('El criterio no es saber la respuesta correcta. El criterio es saber hacer las preguntas correctas antes de actuar. Y para hacer preguntas, no se necesitan años de experiencia, se necesita curiosidad implacable.')
    
    doc.add_page_break()
    
    # Fase 2: Arquitectura Física y Lógica
    doc.add_heading('Fase 2: Arquitectura Física y Lógica', level=2)
    
    doc.add_heading('¿Dónde corre nuestra IA?', level=3)
    doc.add_paragraph('Antes de construir, debemos decidir dónde vivirá nuestro sistema. Esta es una decisión de arquitectura fundamental.')
    
    doc.add_heading('Opción A: Local (Tu propia computadora)', level=4)
    doc.add_paragraph('Analogía: Cocinar en tu casa.')
    
    table5 = doc.add_table(rows=4, cols=2)
    table5.style = 'Table Grid'
    table5.cell(0, 0).text = 'Ventajas'
    table5.cell(0, 1).text = 'Desventajas'
    table5.cell(1, 0).text = 'Gratis'
    table5.cell(1, 1).text = 'Si se daña tu PC, se cae el sistema'
    table5.cell(2, 0).text = 'Tus datos no salen de la PC (máxima privacidad)'
    table5.cell(2, 1).text = 'Si quieres que tu vecino lo use, tiene que ir a tu casa'
    table5.cell(3, 0).text = 'No necesitas internet'
    table5.cell(3, 1).text = 'Limitado por el hardware de tu PC'
    
    doc.add_paragraph('Pensamiento crítico: ¿Manejo datos médicos súper sensibles? → Local.')
    
    doc.add_heading('Opción B: API (Interfaz de Programación de Aplicaciones)', level=4)
    doc.add_paragraph('Analogía: Pedir comida a domicilio por Uber Eats. Tú no cocinas, envías la petición, alguien poderoso cocina, y te devuelve el resultado.')
    
    table6 = doc.add_table(rows=4, cols=2)
    table6.style = 'Table Grid'
    table6.cell(0, 0).text = 'Ventajas'
    table6.cell(0, 1).text = 'Desventajas'
    table6.cell(1, 0).text = 'No necesitas una súper computadora con tarjetas gráficas carísimas'
    table6.cell(1, 1).text = 'Necesitas internet siempre'
    table6.cell(2, 0).text = 'Pagas solo por lo que usas'
    table6.cell(2, 1).text = 'Cada vez que mandas un dato, sale de tu empresa'
    table6.cell(3, 0).text = 'Escalable automáticamente'
    table6.cell(3, 1).text = 'Dependes de un tercero'
    
    doc.add_paragraph('Pensamiento crítico: ¿Mi aplicación va a tener miles de usuarios y no quiero comprar servidores de $10,000? → API.')
    
    doc.add_heading('Opción C: VPS (Servidor Privado Virtual) o Cloud', level=4)
    doc.add_paragraph('Analogía: Alquilar un local comercial. Es tuyo, está encendido 24/7, cualquiera puede entrar por la puerta principal.')
    
    table7 = doc.add_table(rows=4, cols=2)
    table7.style = 'Table Grid'
    table7.cell(0, 0).text = 'Ventajas'
    table7.cell(0, 1).text = 'Desventajas'
    table7.cell(1, 0).text = 'Control total'
    table7.cell(1, 1).text = 'Tienes que mantenerlo (actualizarlo, protegerlo)'
    table7.cell(2, 0).text = 'Accesible desde cualquier lugar del mundo'
    table7.cell(2, 1).text = 'Cuesta dinero mensual'
    table7.cell(3, 0).text = 'Personalizable'
    table7.cell(3, 1).text = 'Requiere conocimientos básicos de administración'
    
    doc.add_paragraph('Pensamiento crítico: ¿Voy a conectar mi IA con una página web para que la usen clientes reales a cualquier hora? → VPS.')
    
    doc.add_heading('Lección de Ingeniería', level=3)
    
    p = doc.add_paragraph()
    run = p.add_run('No existe la mejor arquitectura. Existe la arquitectura que mejor resuelve tu problema específico.')
    run.bold = True
    
    # Guardar documento
    output_path = r'H:\git\TAller de IA\para mi\Taller-de-especializaci-n-de-ia\Clase-05-Arquitectura-Estructura-y-Mentalidad-de-Ingenieria-para-IA\Apuntes_Clase_5_Arquitectura_IA.docx'
    doc.save(output_path)
    print(f'Documento guardado en: {output_path}')

if __name__ == '__main__':
    create_word_document()